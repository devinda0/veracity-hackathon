from __future__ import annotations

from datetime import UTC, datetime
from types import SimpleNamespace
from typing import Any

from bson import ObjectId
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
import pytest

from app.core.app import create_app
from app.core.security import create_access_token, hash_password
from app.db.mongo import MongoDBClient
from app.services.document_parser import DocumentParseError, DocumentParser


class FakeUsersCollection:
    def __init__(self) -> None:
        self.documents: list[dict[str, Any]] = []

    async def find_one(self, query: dict[str, Any]) -> dict[str, Any] | None:
        for document in self.documents:
            if all(document.get(key) == value for key, value in query.items()):
                return document.copy()
        return None


class FakeBusinessContextCollection:
    def __init__(self) -> None:
        self.documents: list[dict[str, Any]] = []

    async def insert_one(self, document: dict[str, Any]) -> SimpleNamespace:
        stored_document = document.copy()
        stored_document["_id"] = ObjectId()
        self.documents.append(stored_document)
        return SimpleNamespace(inserted_id=stored_document["_id"])


class FakeDatabase:
    def __init__(self) -> None:
        self.users = FakeUsersCollection()
        self.business_context = FakeBusinessContextCollection()


@pytest.fixture
def fake_db() -> FakeDatabase:
    original_db = MongoDBClient._db
    database = FakeDatabase()
    MongoDBClient._db = database
    try:
        yield database
    finally:
        MongoDBClient._db = original_db


@pytest.fixture
def client(fake_db: FakeDatabase) -> TestClient:
    with TestClient(create_app()) as test_client:
        yield test_client


def add_user(
    fake_db: FakeDatabase,
    *,
    email: str = "user@example.com",
    password: str = "Secret123!",
    name: str = "Test User",
    is_active: bool = True,
) -> dict[str, Any]:
    user = {
        "_id": ObjectId(),
        "email": email,
        "password_hash": hash_password(password),
        "name": name,
        "created_at": datetime.now(UTC),
        "updated_at": datetime.now(UTC),
        "is_active": is_active,
    }
    fake_db.users.documents.append(user)
    return user


def build_simple_pdf(text: str, *, title: str, author: str) -> bytes:
    def pdf_escape(value: str) -> str:
        return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        (
            f"<< /Length {len(f'BT\\n/F1 18 Tf\\n36 100 Td\\n({pdf_escape(text)}) Tj\\nET'.encode('utf-8'))} >>\n"
            "stream\n"
            f"BT\n/F1 18 Tf\n36 100 Td\n({pdf_escape(text)}) Tj\nET\n"
            "endstream"
        ),
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        f"<< /Title ({pdf_escape(title)}) /Author ({pdf_escape(author)}) /CreationDate (D:20260315090000Z) >>",
    ]

    content = ["%PDF-1.4\n"]
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(sum(len(part.encode("utf-8")) for part in content))
        content.append(f"{index} 0 obj\n{obj}\nendobj\n")

    xref_offset = sum(len(part.encode("utf-8")) for part in content)
    content.append(f"xref\n0 {len(objects) + 1}\n")
    content.append("0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.append(f"{offset:010d} 00000 n \n")
    content.append(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R /Info 6 0 R >>\n")
    content.append(f"startxref\n{xref_offset}\n%%EOF\n")
    return "".join(content).encode("utf-8")


def test_parse_txt_extracts_metadata_and_chunks(tmp_path) -> None:
    text = "Intro paragraph.\n\n" + " ".join(f"token{i}" for i in range(650))
    path = tmp_path / "notes.txt"
    path.write_text(text, encoding="utf-8")

    parsed = DocumentParser.parse_file(path, "txt")

    assert parsed["source"] == "txt"
    assert parsed["metadata"]["title"] == "notes"
    assert parsed["metadata"]["size_bytes"] > 0
    assert "Intro paragraph." in parsed["content"]
    assert len(parsed["chunks"]) >= 2


def test_parse_docx_extracts_text_and_metadata(tmp_path) -> None:
    path = tmp_path / "brief.docx"
    document = DocxDocument()
    document.core_properties.title = "Strategy Brief"
    document.core_properties.author = "Analyst"
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph with detail.")
    document.save(path)

    parsed = DocumentParser.parse_file(path, "docx")

    assert parsed["source"] == "docx"
    assert parsed["metadata"]["title"] == "Strategy Brief"
    assert parsed["metadata"]["author"] == "Analyst"
    assert parsed["metadata"]["paragraphs"] == 2
    assert "Second paragraph with detail." in parsed["content"]


def test_parse_pdf_extracts_text_and_metadata(tmp_path) -> None:
    path = tmp_path / "sample.pdf"
    path.write_bytes(build_simple_pdf("Hello PDF World", title="Quarterly Report", author="Verifier"))

    parsed = DocumentParser.parse_file(path, "pdf")

    assert parsed["source"] == "pdf"
    assert parsed["metadata"]["title"] == "Quarterly Report"
    assert parsed["metadata"]["author"] == "Verifier"
    assert parsed["metadata"]["pages"] == 1
    assert "2026-03-15T09:00:00+00:00" == parsed["metadata"]["date"]
    assert "Hello PDF World" in parsed["content"]


def test_parse_pdf_rejects_corrupted_file(tmp_path) -> None:
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"not-a-real-pdf")

    with pytest.raises(DocumentParseError, match="Unable to parse PDF file"):
        DocumentParser.parse_file(path, "pdf")


def test_upload_document_persists_parsed_content(client: TestClient, fake_db: FakeDatabase) -> None:
    user = add_user(fake_db, email="founder@example.com", name="Founder")
    token = create_access_token({"sub": str(user["_id"]), "email": user["email"]})

    response = client.post(
        "/api/v1/documents/upload-document",
        params={"session_id": "session-123"},
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("context.txt", b"Useful context for the embedding pipeline.", "text/plain")},
    )

    assert response.status_code == 201
    payload = response.json()
    assert payload["filename"] == "context.txt"
    assert payload["chunks"] == 1
    assert payload["metadata"]["title"] == "context"

    stored = fake_db.business_context.documents[0]
    assert stored["user_id"] == str(user["_id"])
    assert stored["session_id"] == "session-123"
    assert stored["source_type"] == "document"
    assert stored["filename"] == "context.txt"
    assert stored["file_type"] == "txt"
    assert stored["content"] == "Useful context for the embedding pipeline."
    assert stored["chunks"] == ["Useful context for the embedding pipeline."]


def test_upload_document_rejects_unsupported_file_type(client: TestClient, fake_db: FakeDatabase) -> None:
    user = add_user(fake_db)
    token = create_access_token({"sub": str(user["_id"]), "email": user["email"]})

    response = client.post(
        "/api/v1/documents/upload-document",
        params={"session_id": "session-123"},
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("context.csv", b"header,value", "text/csv")},
    )

    assert response.status_code == 400
    assert response.json() == {"detail": "File type must be one of: ['docx', 'pdf', 'txt']"}
